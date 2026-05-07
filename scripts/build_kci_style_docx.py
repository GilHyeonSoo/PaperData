from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt


ROOT = Path(__file__).resolve().parents[1]
INPUT_MD = ROOT / "paper" / "draft_kci_style.md"
OUTPUT_DOCX = ROOT / "output" / "doc" / "evtol_paper_kci_style.docx"


def clean_text(text: str) -> str:
    text = text.strip()
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    text = re.sub(r"\*(.*?)\*", r"\1", text)
    return text


def split_caption(text: str) -> tuple[str, str]:
    if "|" in text:
        ko, en = text.split("|", 1)
        return ko.strip(), en.strip()
    return text.strip(), ""


def split_table_row(line: str) -> list[str]:
    line = line.strip()
    if line.startswith("|"):
        line = line[1:]
    if line.endswith("|"):
        line = line[:-1]
    return [cell.strip() for cell in line.split("|")]


def is_table_separator(line: str) -> bool:
    cells = split_table_row(line)
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell.strip()) for cell in cells)


def resolve_image_path(markdown_path: str) -> Path:
    path = Path(markdown_path)
    if path.is_absolute():
        return path
    return (INPUT_MD.parent / path).resolve()


def parse_markdown(markdown_text: str) -> dict[str, object]:
    data: dict[str, object] = {
        "title": "",
        "subtitle": "",
        "english_title": "",
        "summary": "",
        "abstract": "",
        "keywords": "",
        "body": [],
        "references": [],
    }
    mode = "front"
    bucket: list[str] = []

    def flush() -> None:
        nonlocal bucket, mode
        text = "\n".join(bucket).strip()
        if mode in {"english_title", "summary", "abstract", "keywords"}:
            data[mode] = text
        bucket = []

    for raw in markdown_text.splitlines():
        stripped = raw.strip()
        if stripped.startswith("# "):
            data["title"] = stripped[2:].strip()
            continue
        if mode == "front" and stripped and not stripped.startswith("## ") and not data["subtitle"]:
            data["subtitle"] = stripped
            continue
        if stripped == "## English Title":
            flush()
            mode = "english_title"
            continue
        if stripped == "## 요약":
            flush()
            mode = "summary"
            continue
        if stripped == "## Abstract":
            flush()
            mode = "abstract"
            continue
        if stripped == "## Keywords":
            flush()
            mode = "keywords"
            continue
        if stripped.startswith("## I."):
            flush()
            mode = "body"
            data["body"].append(raw)
            continue
        if stripped == "## References":
            flush()
            mode = "references"
            continue
        if mode == "body":
            data["body"].append(raw)
        elif mode == "references":
            data["references"].append(raw)
        elif mode in {"english_title", "summary", "abstract", "keywords"}:
            bucket.append(raw)
    flush()
    return data


def set_run_font(
    run,
    east_asia: str = "휴먼명조",
    latin: str | None = None,
    size: float | None = None,
    bold: bool = False,
) -> None:
    latin = latin or east_asia
    run.font.name = latin
    if size is not None:
        run.font.size = Pt(size)
    run.bold = bold
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), east_asia)
    rfonts.set(qn("w:ascii"), latin)
    rfonts.set(qn("w:hAnsi"), latin)


def set_page(section) -> None:
    section.page_width = Mm(190)
    section.page_height = Mm(260)
    section.top_margin = Mm(17.5)
    section.bottom_margin = Mm(17.5)
    section.left_margin = Mm(20)
    section.right_margin = Mm(20)
    section.header_distance = Mm(10)
    section.footer_distance = Mm(0)


def set_columns(section, count: int = 1, space_twips: int = 340) -> None:
    sect_pr = section._sectPr
    cols = sect_pr.xpath("./w:cols")
    if cols:
        cols_el = cols[0]
    else:
        cols_el = OxmlElement("w:cols")
        sect_pr.append(cols_el)
    cols_el.set(qn("w:num"), str(count))
    cols_el.set(qn("w:space"), str(space_twips))


def set_header(section, title: str) -> None:
    header = section.header
    header.is_linked_to_previous = False
    paragraph = header.paragraphs[0]
    paragraph.text = ""
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("Journal of KIIT. Vol. 23, No. 0, pp. 00-00, 00 (0), 2025. pISSN 1598-8619, eISSN 2093-7571")
    set_run_font(run, east_asia="Times New Roman", latin="Times New Roman", size=8)


def add_horizontal_line(doc: Document) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("_" * 92)
    set_run_font(run, east_asia="Times New Roman", latin="Times New Roman", size=8)


def add_paragraph(
    doc: Document,
    text: str,
    kind: str = "body",
    align: WD_ALIGN_PARAGRAPH | None = None,
) -> None:
    paragraph = doc.add_paragraph()
    if align is not None:
        paragraph.alignment = align
    elif kind in {"title", "subtitle", "english_title", "front_heading", "section", "caption", "caption_en", "refs_heading"}:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    fmt = paragraph.paragraph_format
    fmt.space_after = Pt(0)
    fmt.space_before = Pt(0)
    if kind in {"body", "summary", "abstract"}:
        fmt.first_line_indent = Pt(10)
    if kind == "reference":
        fmt.left_indent = Pt(15)
        fmt.first_line_indent = Pt(-15)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        fmt.line_spacing = 1.5
    elif kind in {"table_cell", "caption", "caption_en"}:
        fmt.line_spacing = 1.3
    else:
        fmt.line_spacing = 1.5

    run = paragraph.add_run(clean_text(text))
    if kind == "title":
        set_run_font(run, east_asia="휴먼명조", size=17, bold=True)
    elif kind == "subtitle":
        set_run_font(run, east_asia="휴먼명조", size=10)
    elif kind == "english_title":
        set_run_font(run, east_asia="Times New Roman", latin="Times New Roman", size=15, bold=True)
    elif kind == "front_heading":
        set_run_font(run, east_asia="휴먼명조", size=9.2)
    elif kind == "summary":
        set_run_font(run, east_asia="휴먼명조", size=9.2)
    elif kind == "abstract":
        set_run_font(run, east_asia="Times New Roman", latin="Times New Roman", size=9.2)
    elif kind == "keywords":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_run_font(run, east_asia="Times New Roman", latin="Times New Roman", size=9.2)
    elif kind == "section":
        set_run_font(run, east_asia="휴먼명조", size=11)
    elif kind == "subsection":
        set_run_font(run, east_asia="휴먼명조", size=10, bold=True)
    elif kind == "caption":
        set_run_font(run, east_asia="휴먼명조", size=8.2)
    elif kind == "caption_en":
        set_run_font(run, east_asia="Times New Roman", latin="Times New Roman", size=8)
    elif kind == "refs_heading":
        set_run_font(run, east_asia="Times New Roman", latin="Times New Roman", size=11)
    elif kind == "reference":
        set_run_font(run, east_asia="Times New Roman", latin="Times New Roman", size=8.5)
    else:
        set_run_font(run, east_asia="휴먼명조", size=9.6)


def flush_body(doc: Document, buffer: list[str]) -> None:
    if not buffer:
        return
    text = " ".join(part.strip() for part in buffer if part.strip())
    buffer.clear()
    if text:
        add_paragraph(doc, text, "body")


def append_table(doc: Document, lines: list[str], caption: str, table_no: int) -> int:
    rows = [split_table_row(line) for line in lines if not is_table_separator(line)]
    if not rows:
        return table_no
    max_cols = max(len(row) for row in rows)
    for row in rows:
        row.extend([""] * (max_cols - len(row)))
    ko_caption, en_caption = split_caption(caption)
    add_paragraph(doc, f"표 {table_no}. {ko_caption}", "caption")
    if en_caption:
        add_paragraph(doc, f"Table {table_no}. {en_caption}", "caption_en")

    table = doc.add_table(rows=len(rows), cols=max_cols)
    table.style = "Table Grid"
    table.autofit = True
    for r_idx, row in enumerate(rows):
        for c_idx, cell_text in enumerate(row):
            cell = table.rows[r_idx].cells[c_idx]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.line_spacing = 1.1
            run = paragraph.add_run(clean_text(cell_text))
            set_run_font(
                run,
                east_asia="Times New Roman",
                latin="Times New Roman",
                size=6.3,
                bold=r_idx == 0,
            )
    doc.add_paragraph()
    return table_no + 1


def append_image(doc: Document, line: str, fig_no: int) -> int:
    match = re.match(r"!\[(.*?)\]\((.*?)\)", line.strip())
    if not match:
        return fig_no
    caption, path_text = match.groups()
    ko_caption, en_caption = split_caption(caption)
    image_path = resolve_image_path(path_text)
    if not image_path.exists():
        add_paragraph(doc, f"[Missing figure: {path_text}]", "caption")
        return fig_no
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Mm(68))
    add_paragraph(doc, f"그림 {fig_no}. {ko_caption}", "caption")
    if en_caption:
        add_paragraph(doc, f"Fig. {fig_no}. {en_caption}", "caption_en")
    return fig_no + 1


def append_body(doc: Document, body_lines: list[str], references: list[str]) -> None:
    buffer: list[str] = []
    table_buffer: list[str] = []
    table_caption = ""
    table_no = 1
    fig_no = 1

    for raw in body_lines:
        stripped = raw.strip()
        if table_buffer and not stripped.startswith("|"):
            table_no = append_table(doc, table_buffer, table_caption, table_no)
            table_buffer = []
            table_caption = ""
        if not stripped:
            flush_body(doc, buffer)
        elif stripped.startswith("## "):
            flush_body(doc, buffer)
            add_paragraph(doc, stripped[3:], "section")
        elif stripped.startswith("### "):
            flush_body(doc, buffer)
            add_paragraph(doc, stripped[4:], "subsection", align=WD_ALIGN_PARAGRAPH.LEFT)
        elif stripped.startswith("[표]"):
            flush_body(doc, buffer)
            table_caption = stripped.replace("[표]", "", 1).strip()
        elif stripped.startswith("|"):
            table_buffer.append(stripped)
        elif stripped.startswith("!["):
            flush_body(doc, buffer)
            fig_no = append_image(doc, stripped, fig_no)
        else:
            buffer.append(stripped)

    if table_buffer:
        append_table(doc, table_buffer, table_caption, table_no)
    flush_body(doc, buffer)
    add_paragraph(doc, "References", "refs_heading")

    ref_buffer: list[str] = []
    for raw in references:
        stripped = raw.strip()
        if not stripped:
            if ref_buffer:
                add_paragraph(doc, " ".join(ref_buffer), "reference")
                ref_buffer = []
        else:
            ref_buffer.append(stripped)
    if ref_buffer:
        add_paragraph(doc, " ".join(ref_buffer), "reference")


def build_docx() -> None:
    OUTPUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    data = parse_markdown(INPUT_MD.read_text(encoding="utf-8"))

    doc = Document()
    first = doc.sections[0]
    set_page(first)
    set_columns(first, 1)
    set_header(first, str(data["title"]))

    add_paragraph(doc, str(data["title"]), "title")
    add_paragraph(doc, str(data["subtitle"]), "subtitle")
    add_paragraph(doc, str(data["english_title"]), "english_title")
    add_horizontal_line(doc)
    add_paragraph(doc, "요  약", "front_heading")
    add_paragraph(doc, str(data["summary"]), "summary")
    add_paragraph(doc, "Abstract", "front_heading")
    add_paragraph(doc, str(data["abstract"]), "abstract")
    add_paragraph(doc, "Keywords: " + str(data["keywords"]), "keywords")
    add_horizontal_line(doc)

    doc.add_page_break()
    body_section = doc.add_section(WD_SECTION.CONTINUOUS)
    set_page(body_section)
    set_columns(body_section, 2)
    set_header(body_section, str(data["title"]))
    append_body(doc, data["body"], data["references"])

    doc.save(OUTPUT_DOCX)
    print(OUTPUT_DOCX)


if __name__ == "__main__":
    build_docx()
